from __future__ import annotations
import io
import json
import re
import tempfile
from typing import Dict, Tuple
import streamlit as st
from bs4 import BeautifulSoup, NavigableString, Tag
from anthropic import Anthropic
# --- опциональные импорты для офисных файлов ---
try:
    from docx import Document # .docx
except Exception:
    Document = None
try:
    import textract # .doc (если установлен и есть системные зависимости)
except Exception:
    textract = None
# ----------------------------
# Ключ и модель из Streamlit Secrets / окружения
# ----------------------------
import os
API_KEY = st.secrets.get("ANTHROPIC_API_KEY") or os.getenv("ANTHROPIC_API_KEY", "")
MODEL_ID = st.secrets.get("ANTHROPIC_MODEL", os.getenv("ANTHROPIC_MODEL", "claude-3-5-sonnet-20240620"))
# ----------------------------
# Оформление
# ----------------------------
st.set_page_config(page_title="Humanizer — сохранение структуры", page_icon="🛠️", layout="wide")
st.title("🛠️ Humanizer с сохранением структуры")
# ----------------------------
# Хелперы
# ----------------------------
def is_html(text: str) -> bool:
    if not text:
        return False
    has_tag = bool(re.search(r"<([a-zA-Z][^>]*?)>", text))
    has_angle = "</" in text or "/>" in text
    return has_tag and has_angle
def extract_text_nodes_as_mapping(html: str) -> Tuple[str, Dict[str, str]]:
    """Оборачивает текстовые узлы в <span data-hid="..."> и собирает mapping id->text."""
    soup = BeautifulSoup(html, "lxml")
    for bad in soup(["script", "style", "noscript"]):
        bad.extract()
    hid_counter = 0
    mapping: Dict[str, str] = {}
    def tag_text_nodes(t: Tag) -> None:
        nonlocal hid_counter
        for child in list(t.children):
            if isinstance(child, NavigableString):
                text = str(child)
                if text and text.strip():
                    hid_counter += 1
                    hid = f"t{hid_counter}"
                    span = soup.new_tag("span")
                    span["data-hid"] = hid
                    span.string = text
                    child.replace_with(span)
                    mapping[hid] = text
            elif isinstance(child, Tag):
                tag_text_nodes(child)
    tag_text_nodes(soup.body or soup)
    return str(soup), mapping
def replace_text_nodes_from_mapping(html_with_ids: str, replacements: Dict[str, str]) -> str:
    soup = BeautifulSoup(html_with_ids, "lxml")
    for span in soup.find_all(attrs={"data-hid": True}):
        hid = span.get("data-hid")
        if hid in replacements:
            span.string = replacements[hid]
    for span in soup.find_all(attrs={"data-hid": True}):
        del span["data-hid"]
    return str(soup)
def _safe_json_loads(maybe_json: str) -> Dict[str, str]:
    """Парсит JSON-объект из ответа модели. Пытается найти первый {...} блок."""
    try:
        return json.loads(maybe_json)
    except Exception:
        pass
    m = re.search(r"\{.*\}", maybe_json, flags=re.DOTALL)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            pass
    raise ValueError("Не удалось распарсить JSON из ответа модели.")
def _word_count(s: str) -> int:
    tokens = re.findall(r"\w+", s, flags=re.UNICODE)
    return len(tokens)
def append_words_marker_to_html(html: str, n: int) -> str:
    """Добавляет в конец HTML видимый маркер [Words: N] как последний <p>."""
    try:
        soup = BeautifulSoup(html, "lxml")
        container = soup.body or soup
        p = soup.new_tag("p")
        p.string = f"[Words: {n}]"
        container.append(p)
        return str(soup)
    except Exception:
        return f"{html}\n[Words: {n}]"
# ----------------------------
# Промпты
# ----------------------------
# 1) Обычный текст → отредактированный текст (+[Words: N]).
PROMPT_PLAIN_TEXT = """Task: Edit the text so it reads naturally and fluently for a native speaker while preserving the original meaning, structure, and tone.
Language: Use the SAME language as the input text (auto-detect). Do NOT translate. Preserve the original dialect/orthography (e.g., en-GB vs en-US).
Requirements:
- Word count: keep within ±2% of the original. Append the final count as [Words: N].
- Structure: keep the same paragraph breaks, headings, list order/numbering, and overall section order.
- Formatting: keep punctuation, quotation marks, inline formatting (bold/italic/links/code), emojis, citation markers, and references exactly as they are.
- Facts & entities: do not add, remove, or alter information. Keep names, numbers, dates, URLs, and titles unchanged.
- Tone & register: preserve the author’s voice, level of formality, and rhetorical stance.
- Style tweaks: replace awkward phrasing with idiomatic expressions, reduce repetitiveness, vary sentence length for burstiness (mix short and long sentences), and simplify clunky constructions—without changing emphasis or intent. Increase perplexity by using unexpected but natural phrasing, contractions (e.g., don't instead of do not), colloquialisms, and varied vocabulary to avoid predictable patterns.
- Non-text elements (code, formulas, tables): leave unchanged.
- If perfect word-count preservation would hurt clarity or grammar, prefer clarity but stay as close as possible to the target range.
- Return ONLY the edited text—no explanations, no metadata (besides [Words: N]), no code fences.
Input:
"""
# 2) Обычный текст → красивый семантический HTML И тоже добавить [Words: N] как последний <p>.
PROMPT_PLAIN_TO_HTML = """Task: Edit the text so it reads naturally and fluently for a native speaker while preserving the original meaning, structure, and tone — then output clean, semantic HTML for the edited text.
Language: Use the SAME language as the input text (auto-detect). Do NOT translate. Preserve the original dialect/orthography.
Requirements:
- Word count: keep within ±2% of the original. Append the final count as a visible last paragraph: [Words: N].
- Keep the same paragraph breaks, headings, list order/numbering, and overall section order as in the input. Convert them to equivalent HTML structure.
- Preserve punctuation, quotation marks, inline emphasis/links/code semantics; convert inline formatting markers to their HTML equivalents (<strong>/<em>/<a>/<code>).
- Do not add, remove, or alter facts, names, numbers, dates, URLs, or titles.
- Preserve the author’s voice and register; improve fluency without changing intent. Vary sentence length for burstiness (mix short and long), increase perplexity with unexpected natural phrasing, use contractions, idioms, and varied vocab to mimic human writing and avoid AI detection patterns.
- If the edited content contains at least one <table>, INCLUDE at the very top of the output a SINGLE <style> block:
  <style>
  table { border-collapse: collapse; }
  table, th, td { border: 1px solid #000; }
  </style>
  Do not include any other CSS or inline style attributes.
- Non-text elements (code blocks, formulas, tables) should be kept as-is but wrapped in appropriate HTML tags (<pre><code>, <table>, etc.) if present.
- Return ONLY the HTML markup of the edited text. No markdown, no comments, no code fences, no explanations.
- Allowed tags include: style (single block as specified above), p, h1..h4, ul/ol/li, blockquote, pre, code, a, strong, em, table/thead/tbody/tr/th/td, img (only if present in input), span.
Input:
"""
# 3) HTML через JSON-замену (сохраняем исходную разметку 1:1); [Words: N] добавим в приложении.
PROMPT_HTML_JSON = """You will be provided with a JSON object containing key-value pairs where keys are IDs and values are text strings extracted from HTML.

Task: Edit each value so it reads naturally and fluently for a native speaker while preserving the original meaning, structure, and tone.

Language: Use the SAME language as each input value (auto-detect). Do NOT translate. Preserve the original dialect/orthography.

Requirements (APPLY PER VALUE):
- Word count: keep within ±2% of that value’s original length. Do NOT add any extra markers like [Words: N].
- Keep punctuation, quotation marks, inline formatting, emojis, citation markers, and references exactly as they are in the value.
- Do not add, remove, or alter facts, names, numbers, dates, URLs, or titles.
- Preserve the author’s voice and register.
- Absolutely do NOT introduce or remove HTML tags; you are editing TEXT CONTENT ONLY.
- Improve fluency by varying sentence length for burstiness (mix short/long), increasing perplexity with unexpected natural phrasing, using contractions, idioms, and varied vocabulary to avoid predictable AI patterns.
- Return ONLY a valid JSON object with the SAME KEYS and improved string values. No comments, no code fences, no extra text, no explanations. Ensure the output is parseable as JSON.
"""
# ----------------------------
# Работа с моделями
# ----------------------------
def call_anthropic_json_map(api_key: str, mapping: Dict[str, str]) -> Dict[str, str]:
    client = Anthropic(api_key=api_key)
    input_json = json.dumps(mapping, ensure_ascii=False)
    resp = client.messages.create(
        model=MODEL_ID,
        max_tokens=8192,  # Увеличьте при необходимости для больших текстов
        system=PROMPT_HTML_JSON,
        messages=[
            {"role": "user", "content": input_json},
        ],
    )
    content = resp.content[0].text if resp.content else "{}"
    return _safe_json_loads(content)
def call_anthropic_rewrite_text(api_key: str, text: str) -> str:
    """Обычный текст → отредактированный текст (+[Words: N])."""
    client = Anthropic(api_key=api_key)
    resp = client.messages.create(
        model=MODEL_ID,
        max_tokens=8192,
        messages=[
            {"role": "user", "content": PROMPT_PLAIN_TEXT + text},
        ],
    )
    return (resp.content[0].text or "").strip()
def call_anthropic_rewrite_text_to_html(api_key: str, text: str) -> str:
    """Обычный текст → красивый семантический HTML с финальным <p>[Words: N]</p>."""
    client = Anthropic(api_key=api_key)
    resp = client.messages.create(
        model=MODEL_ID,
        max_tokens=8192,
        messages=[
            {"role": "user", "content": PROMPT_PLAIN_TO_HTML + text},
        ],
    )
    return (resp.content[0].text or "").strip()
# ----------------------------
# Загрузка файлов
# ----------------------------
def read_text_file(uploaded) -> str:
    raw = uploaded.read().decode("utf-8", errors="ignore")
    return raw
def read_docx_file(uploaded) -> str:
    if Document is None:
        raise RuntimeError("Не установлен пакет python-docx. Установите: pip install python-docx")
    uploaded.seek(0)
    doc = Document(uploaded)
    blocks = []
    for p in doc.paragraphs:
        blocks.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            blocks.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(block for block in blocks if block is not None)
def read_doc_file(uploaded) -> str:
    if textract is None:
        raise RuntimeError("Для чтения .doc установите textract: pip install textract")
    uploaded.seek(0)
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=True) as tmp:
        tmp.write(uploaded.read())
        tmp.flush()
        data = textract.process(tmp.name)
    return data.decode("utf-8", errors="ignore")
# ----------------------------
# Генерация скачиваемых файлов
# ----------------------------
def build_docx_bytes(plain_text: str) -> bytes:
    if Document is None:
        raise RuntimeError("Для экспорта в .docx установите python-docx: pip install python-docx")
    doc = Document()
    for para in plain_text.split("\n"):
        doc.add_paragraph(para)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
# ----------------------------
# UI: ввод и выходной формат
# ----------------------------
col_in, col_opts = st.columns([2, 1], gap="large")
with col_in:
    st.markdown("#### Вставьте текст или HTML")
    input_text = st.text_area(
        "", height=280,
        placeholder="Вставьте сюда ваш текст / HTML. Структура и порядок элементов будут сохранены.",
    )
    uploaded = st.file_uploader(
        "…или загрузите файл (.html, .txt, .md, .docx, .doc)",
        type=["html", "txt", "md", "docx", "doc"],
        accept_multiple_files=False
    )
    if uploaded is not None and not input_text:
        ext = (uploaded.name.split(".")[-1] or "").lower()
        try:
            if ext in {"html", "htm"}:
                input_text = read_text_file(uploaded)
            elif ext in {"txt", "md"}:
                input_text = read_text_file(uploaded)
            elif ext == "docx":
                input_text = read_docx_file(uploaded)
            elif ext == "doc":
                input_text = read_doc_file(uploaded)
            else:
                st.error("Неподдерживаемый формат файла.")
        except Exception as e:
            st.error(f"Не удалось прочитать файл: {e}")
with col_opts:
    st.markdown("#### Выходной формат")
    out_format = st.radio("Формат выдачи", ["HTML", "Plain/Markdown"], index=0, horizontal=True)
    text_download_fmt = st.selectbox("Скачать текст как", ["TXT", "MD", "DOCX"], index=0, help="Применяется, когда результат — текст.")
    st.markdown("#### Обработать")
    go = st.button("🚀 Запустить хуманизацию", type="primary", use_container_width=True)
# ----------------------------
# Основная логика
# ----------------------------
if go:
    if not input_text or not input_text.strip():
        st.error("Пожалуйста, вставьте текст или загрузите файл.")
    elif not API_KEY:
        st.error(
            "Не найден ANTHROPIC_API_KEY. Укажите его в Streamlit secrets "
            "(Settings → Secrets) или в переменной окружения."
        )
    else:
        try:
            with st.spinner("Обработка текста моделью…"):
                if is_html(input_text):
                    # HTML → JSON-замена (сохраняем исходные теги 1:1). [Words: N] добавляем после сборки.
                    html_with_ids, mapping = extract_text_nodes_as_mapping(input_text)
                    rewritten_map = call_anthropic_json_map(API_KEY, mapping)
                    result_html = replace_text_nodes_from_mapping(html_with_ids, rewritten_map)
                    # Считаем слова по видимому тексту и добавляем маркер
                    visible_text = BeautifulSoup(result_html, "lxml").get_text(separator=" ").strip()
                    words_n = _word_count(visible_text)
                    result_html = append_words_marker_to_html(result_html, words_n)
                    if out_format == "HTML":
                        result = result_html
                        out_kind = "html"
                    else:
                        plain = BeautifulSoup(result_html, "lxml").get_text(separator="\n")
                        plain = re.sub(r"\n{3,}", "\n\n", plain).strip()
                        result = plain
                        out_kind = "txt"
                else:
                    # Обычный текст/Markdown
                    if out_format == "HTML":
                        # Красивый HTML — просим модель также добавить [Words: N] как последний параграф.
                        result = call_anthropic_rewrite_text_to_html(API_KEY, input_text)
                        out_kind = "html"
                    else:
                        # Отредактированный текст с [Words: N] в конце (добавляет модель)
                        result = call_anthropic_rewrite_text(API_KEY, input_text)
                        out_kind = "txt"
            # Вывод и скачивание
            st.success("Готово!")
            if out_kind == "html":
                st.markdown("Просмотр HTML:")
                st.components.v1.html(result, height=600, scrolling=True)
                st.download_button(
                    label="⬇️ Скачать .html",
                    data=result.encode("utf-8"),
                    file_name="humanized.html",
                    mime="text/html",
                )
                with st.expander("Показать HTML-код"):
                    st.code(result, language="html")
            else:
                st.markdown("Предпросмотр текста:")
                st.text_area("", value=result, height=400)
                fmt = text_download_fmt.upper()
                if fmt == "TXT":
                    st.download_button(
                        label="⬇️ Скачать .txt",
                        data=result.encode("utf-8"),
                        file_name="humanized.txt",
                        mime="text/plain",
                    )
                elif fmt == "MD":
                    st.download_button(
                        label="⬇️ Скачать .md",
                        data=result.encode("utf-8"),
                        file_name="humanized.md",
                        mime="text/markdown",
                    )
                elif fmt == "DOCX":
                    try:
                        docx_bytes = build_docx_bytes(result)
                        st.download_button(
                            label="⬇️ Скачать .docx",
                            data=docx_bytes,
                            file_name="humanized.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        )
                    except Exception as e:
                        st.error(f"Не удалось сформировать .docx: {e}")
        except Exception as e:
            st.error(f"Ошибка обработки: {e}")
